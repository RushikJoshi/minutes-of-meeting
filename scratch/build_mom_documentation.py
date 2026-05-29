from __future__ import annotations

import os
from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\pc\GitHub\minutes-of-meeting")
OUT_DIR = ROOT / "scratch" / "documentation_assets"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = ROOT / "MOM_Project_Complete_Documentation.docx"


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GRAY = "F2F4F7"
DARK = "1F2937"
MUTED = "4B5563"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D0D7DE", sz="6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill=GRAY) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 41, 55)


def add_table(doc, headers, rows, widths=None, header_fill=GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    style_table(table, header_fill)
    doc.add_paragraph()
    return table


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="B6C7D6")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(body)
    doc.add_paragraph()


def font(size=18):
    candidates = [
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_flow(name, title, lanes, steps, edges, width=1600, height=900):
    image = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(image)
    title_font = font(34)
    lane_font = font(22)
    box_font = font(20)
    small_font = font(16)
    d.rectangle((0, 0, width, 74), fill="#1F4E79")
    d.text((40, 20), title, fill="white", font=title_font)
    lane_w = width / len(lanes)
    for i, lane in enumerate(lanes):
        x0 = int(i * lane_w)
        x1 = int((i + 1) * lane_w)
        fill = "#F6F8FA" if i % 2 == 0 else "#EEF4FA"
        d.rectangle((x0, 74, x1, height), fill=fill, outline="#D0D7DE")
        d.text((x0 + 24, 95), lane, fill="#1F2937", font=lane_font)
    centers = {}
    for step_id, lane_idx, y, text, color in steps:
        x = int(lane_idx * lane_w + 35)
        box_w = int(lane_w - 70)
        box_h = 80
        d.rounded_rectangle((x, y, x + box_w, y + box_h), radius=14, fill=color, outline="#6B7280", width=2)
        words = text.split()
        lines, current = [], ""
        for word in words:
            if len(current + " " + word) > 28:
                lines.append(current.strip())
                current = word
            else:
                current += " " + word
        if current:
            lines.append(current.strip())
        for j, line in enumerate(lines[:3]):
            d.text((x + 18, y + 16 + j * 21), line, fill="#111827", font=box_font)
        centers[step_id] = (x + box_w // 2, y + box_h // 2)
    for src, dst, label in edges:
        sx, sy = centers[src]
        dx, dy = centers[dst]
        d.line((sx, sy, dx, dy), fill="#374151", width=3)
        # arrow head
        if dx >= sx:
            d.polygon([(dx, dy), (dx - 14, dy - 8), (dx - 14, dy + 8)], fill="#374151")
        else:
            d.polygon([(dx, dy), (dx + 14, dy - 8), (dx + 14, dy + 8)], fill="#374151")
        if label:
            lx, ly = (sx + dx) // 2, (sy + dy) // 2
            d.rectangle((lx - 70, ly - 14, lx + 70, ly + 14), fill="white", outline="#CBD5E1")
            d.text((lx - 62, ly - 10), label[:18], fill="#374151", font=small_font)
    path = OUT_DIR / f"{name}.png"
    image.save(path)
    return path


def build_diagrams():
    diagrams = {}
    diagrams["architecture"] = draw_flow(
        "architecture",
        "System Architecture",
        ["React Client", "Express API", "Services", "MongoDB / Files"],
        [
            ("ui", 0, 170, "Protected React routes, pages, forms", "#D9EAF7"),
            ("api", 0, 350, "Axios client adds JWT and workspace header", "#D9EAF7"),
            ("routes", 1, 190, "Express /api/v1 route layer", "#E2F0D9"),
            ("middleware", 1, 360, "Auth + workspace middleware", "#E2F0D9"),
            ("services", 2, 190, "Email, PDF, AI, Calendar, Collab services", "#FFF2CC"),
            ("cron", 2, 390, "Reminder loop and auto-complete jobs", "#FFF2CC"),
            ("db", 3, 180, "MongoDB models and Yjs documents", "#FCE4D6"),
            ("uploads", 3, 390, "Uploads folder for files and PDFs", "#FCE4D6"),
        ],
        [
            ("ui", "api", "user action"),
            ("api", "routes", "HTTP"),
            ("routes", "middleware", "guard"),
            ("middleware", "services", "controller"),
            ("services", "db", "CRUD"),
            ("services", "uploads", "PDF/files"),
            ("cron", "db", "scheduled"),
        ],
    )
    diagrams["auth"] = draw_flow(
        "auth_workspace_flow",
        "Authentication and Workspace Flow",
        ["Browser", "Auth Provider", "API/Middleware", "Database"],
        [
            ("login", 0, 160, "User submits login or register form", "#D9EAF7"),
            ("store", 1, 160, "Store mom.token and set Authorization header", "#E2F0D9"),
            ("me", 1, 330, "Load /auth/me on app start", "#E2F0D9"),
            ("jwt", 2, 220, "requireAuth verifies JWT or API key", "#FFF2CC"),
            ("ws", 2, 410, "requireWorkspace validates x-workspace-id", "#FFF2CC"),
            ("user", 3, 220, "User document without passwordHash", "#FCE4D6"),
            ("member", 3, 430, "Membership provides workspace role", "#FCE4D6"),
        ],
        [("login", "store", "success"), ("store", "me", "reload"), ("me", "jwt", "Bearer"), ("jwt", "user", "find"), ("jwt", "ws", "next"), ("ws", "member", "lookup")],
    )
    diagrams["meeting"] = draw_flow(
        "meeting_lifecycle",
        "Meeting Lifecycle",
        ["Frontend", "Meeting Controller", "External Services", "Persistence"],
        [
            ("form", 0, 150, "CreateMeeting submits MeetingForm payload", "#D9EAF7"),
            ("normalize", 1, 145, "Validate, normalize participants, compute endTime", "#E2F0D9"),
            ("platform", 2, 145, "Google Meet, Zoom, or Teams link creation", "#FFF2CC"),
            ("save", 3, 145, "Meeting saved with participants and graph IDs", "#FCE4D6"),
            ("invite", 2, 330, "Send background invite emails with tokenized join URLs", "#FFF2CC"),
            ("start", 1, 510, "Start meeting: scheduled becomes ongoing", "#E2F0D9"),
            ("end", 1, 670, "End meeting: attendance, MOM, PDF generation", "#E2F0D9"),
            ("files", 3, 650, "MOM, ActionItem, PDF URL updated", "#FCE4D6"),
        ],
        [("form", "normalize", "POST"), ("normalize", "platform", "online"), ("platform", "save", "link"), ("save", "invite", "async"), ("save", "start", "later"), ("start", "end", "complete"), ("end", "files", "write")],
    )
    diagrams["mom"] = draw_flow(
        "mom_editor_flow",
        "MOM Editor Save and Publish Flow",
        ["MinutesEditor", "WordLikeEditor", "MOM Controller", "Models/Services"],
        [
            ("load", 0, 150, "GET meeting/:id/minutes loads meeting and MOM", "#D9EAF7"),
            ("template", 0, 330, "Apply custom/default template placeholders", "#D9EAF7"),
            ("edit", 1, 230, "TipTap editor captures formatted HTML", "#E2F0D9"),
            ("auto", 1, 430, "Autosave/manual save sends contentHtml", "#E2F0D9"),
            ("upsert", 2, 310, "Upsert MOM, increment version, publish when requested", "#FFF2CC"),
            ("extract", 3, 220, "Extract and merge action items from HTML", "#FCE4D6"),
            ("sync", 3, 420, "Sync ActionItem collection and meeting.momContent", "#FCE4D6"),
            ("pdf", 2, 610, "Published manual MOM can trigger end lifecycle", "#FFF2CC"),
        ],
        [("load", "template", "blank"), ("template", "edit", "HTML"), ("edit", "auto", "change"), ("auto", "upsert", "PUT"), ("upsert", "extract", "parse"), ("extract", "sync", "save"), ("upsert", "pdf", "publish")],
    )
    diagrams["data"] = draw_flow(
        "data_model_map",
        "Core Data Model Relationships",
        ["Tenant", "Meeting", "MOM", "Outputs"],
        [
            ("workspace", 0, 155, "Workspace", "#D9EAF7"),
            ("membership", 0, 330, "Membership: user role per workspace", "#D9EAF7"),
            ("meeting", 1, 230, "Meeting: schedule, participants, lifecycle", "#E2F0D9"),
            ("mom", 2, 150, "Mom: rich minutes, sections, status, version", "#FFF2CC"),
            ("action", 2, 360, "ActionItem: extracted task tracker", "#FFF2CC"),
            ("share", 3, 150, "Share: secure token and expiry", "#FCE4D6"),
            ("attach", 3, 330, "Attachment and uploads", "#FCE4D6"),
            ("notify", 3, 510, "Notification reminders", "#FCE4D6"),
        ],
        [("workspace", "meeting", "owns"), ("workspace", "membership", "access"), ("meeting", "mom", "1:1"), ("mom", "action", "sync"), ("meeting", "share", "public link"), ("mom", "attach", "files"), ("meeting", "notify", "reminders")],
    )
    return diagrams


def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, "1F2937", 7, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(10.5)
        styles[name].paragraph_format.space_after = Pt(3)
        styles[name].paragraph_format.line_spacing = 1.1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("MOM Project Documentation | Complete Workflow, Frontend, Backend, Database, and Deployment Notes")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(107, 114, 128)


def add_cover(doc):
    logo = ROOT / "client" / "src" / "assets" / "logo_final.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo.exists():
        p.add_run().add_picture(str(logo), width=Inches(1.25))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Minutes of Meeting (MOM)")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Complete Project Documentation and Workflow Report")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(DARK)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Prepared for meeting presentation | Generated on {date.today().strftime('%d %B %Y')}")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()
    add_callout(
        doc,
        "Document purpose",
        "This report explains the MOM project from end to end: why it exists, how users move through the frontend, how backend APIs process each action, how MongoDB stores the data, and how reminders, collaboration, action items, sharing, PDF export, and integrations work together.",
        fill="EAF2F8",
    )
    add_table(
        doc,
        ["Document Area", "Coverage"],
        [
            ["Frontend", "React routes, protected layout, pages, editor behavior, API client, state providers, and user workflows."],
            ["Backend", "Express server, middleware, controllers, services, route map, database models, scheduled jobs, file/PDF handling, and integrations."],
            ["Data and Flow", "Authentication, workspace selection, meeting lifecycle, MOM editing, action item sync, sharing, reminders, and export flow."],
            ["Presentation Use", "Designed as a professional report with detailed explanations, diagrams, tables, and dense page content."],
        ],
        widths=[1.7, 4.8],
        header_fill=LIGHT_BLUE,
    )
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    sections = [
        "1. Executive Overview",
        "2. Project Objectives and Scope",
        "3. Technology Stack",
        "4. High-Level Architecture",
        "5. File Structure",
        "6. Frontend Application Flow",
        "7. Backend Application Flow",
        "8. Authentication and Workspace Security",
        "9. Meeting Scheduling and Lifecycle",
        "10. MOM Editor and Document Workflow",
        "11. Real-Time Collaboration Design",
        "12. Action Items and Task Tracking",
        "13. Reminders, Notifications, and Cron Jobs",
        "14. PDF Export and Sharing",
        "15. Integrations: Google, Microsoft Teams, Zoom, AI",
        "16. Database Schema",
        "17. API Endpoint Reference",
        "18. Error Handling, Validation, and Security",
        "19. Deployment and Environment Configuration",
        "20. End-to-End Workflow Summary",
        "21. Future Improvements",
    ]
    for item in sections:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item.split(". ", 1)[1] if ". " in item else item)
    add_callout(doc, "Reading note", "The document is written for presentation and project explanation. Diagrams explain the sequence quickly; the paragraphs after each diagram provide implementation-level detail.")
    doc.add_page_break()


def add_image(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.7))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(75, 85, 99)


def file_structure_text():
    include_dirs = ["client/src", "server/controllers", "server/models", "server/routes", "server/services", "server/middlewares", "server/utils"]
    lines = ["minutes-of-meeting/"]
    for base in include_dirs:
        folder = ROOT / base
        if not folder.exists():
            continue
        lines.append(f"  {base}/")
        for path in sorted(folder.rglob("*")):
            if path.is_file():
                rel = path.relative_to(ROOT).as_posix()
                if any(skip in rel for skip in ["node_modules", "uploads"]):
                    continue
                indent = "    " + "  " * (len(path.relative_to(folder).parts) - 1)
                lines.append(f"{indent}{path.name}")
    return "\n".join(lines)


def add_file_structure(doc):
    doc.add_heading("5. File Structure", level=1)
    add_para(doc, "The project is divided into two main applications: the React frontend in client and the Express backend in server. The separation is clean enough for presentation: frontend files manage user interface and browser state, while backend files manage API behavior, business rules, persistence, files, notifications, and external integrations.")
    add_table(
        doc,
        ["Folder", "Responsibility"],
        [
            ["client/src/pages", "Route-level screens such as Dashboard, Meetings, MeetingDetails, MinutesEditor, ActionItems, Settings, Reports, and sharing/join pages."],
            ["client/src/components", "Reusable UI pieces such as MeetingForm, MeetingCard, ParticipantPicker, calendar components, and the WordLikeEditor."],
            ["client/src/providers", "AuthProvider and WorkspaceProvider keep token, current user, workspace list, and active workspace available across the app."],
            ["server/routes", "Maps HTTP endpoints to middleware and controller methods under /api/v1."],
            ["server/controllers", "Receives requests, validates payloads, calls services/models, and returns API responses."],
            ["server/models", "Mongoose schemas for User, Workspace, Membership, Meeting, Mom, ActionItem, Share, Attachment, Notification, IntegrationToken, and API keys."],
            ["server/services", "Business-service layer for email, PDF, reminders, collaboration, AI, Google Calendar, Microsoft Graph, Zoom, and action-item syncing."],
            ["server/utils", "Shared helpers for DB connection, tokens, PDF templates, mailer, public client URL, Chrome path, async handling, and errors."],
        ],
        widths=[1.7, 4.8],
        header_fill=LIGHT_BLUE,
    )
    p = doc.add_paragraph()
    r = p.add_run(file_structure_text())
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(7.4)
    doc.add_page_break()


def add_sections(doc, diagrams):
    doc.add_heading("1. Executive Overview", level=1)
    add_para(doc, "The Minutes of Meeting project is a full-stack web application for planning meetings, inviting participants, recording minutes, generating action items, sending reminders, and exporting or sharing finalized MOM documents. It combines a React frontend with an Express and MongoDB backend. The application is built around a workspace-aware model so that each user can operate inside an organization or project space without mixing data with other workspaces.")
    add_para(doc, "The core business problem solved by the system is that meeting information is usually scattered across calendar invites, chat messages, notes, manual documents, emails, and follow-up trackers. This project brings those activities into one flow: create the meeting, invite participants, conduct the meeting, write the minutes, extract tasks, publish the MOM, generate a PDF, and keep follow-up items visible until they are completed.")
    add_callout(doc, "Project summary", "MOM acts as a meeting operating system. The frontend gives users an organized dashboard, meeting screens, calendar view, and Word-like editor. The backend enforces authentication and workspace access, persists records in MongoDB, sends reminders and invitations, integrates with calendar/video platforms, and generates PDFs.")
    add_bullets(doc, [
        "Users authenticate through JWT-based login/register flows and the frontend stores the token in localStorage under mom.token.",
        "Every protected request carries both Authorization and x-workspace-id headers so the backend can identify the user and workspace context.",
        "Meetings support title, agenda, date, start/end time, type, platform, location, priority, participants, reminders, and lifecycle status.",
        "Minutes are stored as structured MOM fields plus rich HTML content, versioned and linked one-to-one with a meeting.",
        "Action items can be manually entered or extracted from MOM HTML, then synchronized into a dedicated ActionItem collection.",
        "The project includes PDF generation, secure share links, notifications, attachments, reports, and external integration hooks.",
    ])

    doc.add_heading("2. Project Objectives and Scope", level=1)
    add_para(doc, "The objective is not only to create a note-taking page, but to manage the entire lifecycle of a meeting. A user should be able to start from a blank dashboard, create a meeting, invite internal or external participants, receive reminders, open the MOM editor, write or publish the minutes, and then track action items after the meeting. For a presentation, this scope is important because it shows the project as a workflow system rather than a single CRUD application.")
    add_table(doc, ["Objective", "Implementation in Project"], [
        ["Centralized meeting scheduling", "CreateMeeting and MeetingForm call POST /create-meeting; Meeting model stores schedule, participants, platform, and status."],
        ["Secure access", "ProtectedRoute, AuthProvider, requireAuth, requireWorkspace, and Membership enforce user and workspace context."],
        ["Professional MOM creation", "MinutesEditor uses WordLikeEditor with TipTap, templates, autosave, publish status, attachments, and PDF export."],
        ["Task accountability", "MOM actionItems are synchronized to ActionItem so employees can view pending/completed work separately."],
        ["Post-meeting deliverables", "PDF generation and share tokens allow MOM output to be distributed after completion."],
        ["Automation", "Reminder service runs cron jobs for meeting reminders, action-item due alerts, and automatic meeting completion."],
    ], widths=[2.0, 4.5])

    doc.add_heading("3. Technology Stack", level=1)
    add_table(doc, ["Layer", "Technologies", "Why It Is Used"], [
        ["Frontend", "React 19, Vite, React Router DOM, Tailwind CSS", "Fast SPA development, route-based screens, protected layouts, and responsive styling."],
        ["Data Fetching", "Axios, TanStack React Query", "Central API client, caching, mutation handling, invalidation after saves/updates, and clean loading/error states."],
        ["Editor", "TipTap, ProseMirror extensions, Yjs-related packages", "Word-like rich text, headings, formatting, links, tables, images, and real-time collaboration foundation."],
        ["Backend", "Node.js, Express 5", "REST API layer, middleware, controllers, static uploads, and service orchestration."],
        ["Database", "MongoDB, Mongoose", "Flexible documents for meetings/MOM content, indexed workspace queries, and schema-based validation."],
        ["Auth/Security", "JWT, bcryptjs, API key hashing", "Session auth for web users and API key access for third-party integration style usage."],
        ["Automation", "node-cron", "Runs reminder loop, due task notifications, and automatic completion of meetings."],
        ["Export/Email", "Puppeteer, Nodemailer", "Generates PDF output and sends invite/reminder emails."],
        ["Integrations", "Google APIs, Microsoft MSAL/Graph, Zoom service", "Creates/syncs online meetings and calendar events where credentials are configured."],
        ["Collaboration", "Hocuspocus Server, y-mongodb/Yjs packages", "Stores collaborative document state and supports editor synchronization through a separate collab server."],
    ], widths=[1.2, 2.0, 3.3], header_fill=LIGHT_BLUE)

    doc.add_heading("4. High-Level Architecture", level=1)
    add_image(doc, diagrams["architecture"], "Figure 1: Main architecture showing React client, Express API, services, MongoDB, and file outputs.")
    add_para(doc, "The React client is responsible for user interaction, form state, navigation, visual layout, and editor behavior. It does not directly access MongoDB or server-side files. Instead, it communicates with the backend through the Axios instance in client/src/api/api.js. This API client computes the backend base URL, prefixes routes with /api/v1, attaches the active workspace header, and clears stale workspace state if the backend rejects access.")
    add_para(doc, "The Express backend acts as the system boundary. server.js mounts all major route modules under /api/v1, serves uploaded files under /uploads, connects to MongoDB, starts the reminder loop, and starts the Hocuspocus collaboration server. Controllers remain focused on request/response behavior, while service files contain heavier operations such as PDF rendering, email delivery, action-item sync, AI processing, reminders, calendar sync, and collaboration persistence.")
    add_para(doc, "MongoDB is the single source of truth for structured application data. It stores users, workspaces, membership roles, meeting records, MOM documents, extracted action items, notifications, attachments, integration tokens, API keys, share tokens, and Yjs collaboration data. The uploads folder stores generated PDFs and uploaded images/files that are referenced from database records.")
    doc.add_page_break()

    add_file_structure(doc)

    doc.add_heading("6. Frontend Application Flow", level=1)
    add_para(doc, "The frontend begins in main.jsx, where React renders the application with global providers. App.jsx then defines all public and protected routes. Public routes include login, register, role selection, document verification, share view, and join meeting. Protected routes are nested under ProtectedRoute and AppLayout, giving authenticated users access to dashboard, admin dashboard, employee dashboard, calendar, meetings, meeting details, MOM editor, settings, profile, documents, reports, action items, and template builder.")
    add_table(doc, ["Frontend Area", "Important Files", "Detailed Role"], [
        ["Routing", "App.jsx", "Defines public and protected route tree. Redirects / to /dashboard and nests workspace screens inside AppLayout."],
        ["Authentication state", "providers/AuthProvider.jsx", "Loads token from localStorage, calls /auth/me, manages login/register/logout, and clears state on 401 responses."],
        ["Workspace state", "providers/WorkspaceProvider.jsx", "Loads workspaces, stores active workspace in mom.workspaceId, and updates API request header."],
        ["API layer", "api/api.js", "Creates Axios instance with /api/v1 base URL, Authorization header helper, workspace header interceptor, and stale workspace cleanup."],
        ["Meeting screens", "CreateMeeting, Meetings, MeetingDetails, Calendar", "Create, list, filter, view, update, start, end, invite, and navigate into MOM editor."],
        ["MOM authoring", "MinutesEditor + WordLikeEditor", "Loads meeting/MOM data, applies templates, autosaves rich HTML, detects task sentences, uploads images, publishes and exports PDF."],
        ["Task tracking", "ActionItems.jsx", "Fetches /action-items with filters and updates task status through PATCH /action-items/:id/status."],
        ["Settings", "Settings.jsx", "Controls Google/Microsoft integration status, connect/disconnect/sync, preferences, and API keys."],
    ], widths=[1.55, 1.65, 3.3])
    add_para(doc, "A typical authenticated page does not manually manage token or workspace headers. Instead, AuthProvider and WorkspaceProvider maintain those concerns globally. When a user logs in, the token is stored and set as a Bearer token. When a workspace is selected, its id is stored and attached to every request. This keeps route components clean: pages call API.get or API.post, and the shared client adds the correct security context.")
    add_para(doc, "React Query is used for most data screens so that changes are reflected across the application. For example, after saving MOM content, MinutesEditor invalidates meeting-minutes, meeting, action-items, and meetings queries. This ensures the editor, meeting detail page, dashboard previews, and action tracker all receive fresh data instead of showing stale information.")
    add_callout(doc, "Frontend user experience", "The frontend is workflow-first: dashboard and lists provide overview, detail pages provide control actions, the editor provides document authoring, and settings provides system integrations. This is better for presentation than explaining files in isolation because every screen has a clear business purpose.")

    doc.add_heading("7. Backend Application Flow", level=1)
    add_para(doc, "The backend entry point is server/server.js. It creates an Express app, enables JSON parsing and CORS, imports route modules, mounts them through an apiRouter under /api/v1, serves static uploads, installs notFound and errorHandler middleware, loads configuration, connects to MongoDB, starts the HTTP server, starts the reminder loop, and attempts to start the Hocuspocus collaboration server.")
    add_table(doc, ["Backend Layer", "Files", "Responsibility"], [
        ["Entry point", "server.js", "Application wiring, route mounting, database connection, reminder loop, collab server startup, uploads static serving."],
        ["Routes", "server/routes/*.js", "Defines endpoint URLs, HTTP methods, and middleware chain before controller functions."],
        ["Controllers", "server/controllers/*.js", "Validates request data, queries or updates models, calls services, and sends API responses."],
        ["Middlewares", "authMiddleware, workspaceMiddleware, roleMiddleware, validateMiddleware", "Authenticates users/API keys, resolves workspace membership, checks roles, and handles validation."],
        ["Services", "server/services/*.js", "Encapsulates non-trivial work: email, reminders, PDF, AI, collaboration, integrations, and action item sync."],
        ["Models", "server/models/*.js", "Mongoose schema definitions and indexes for persistence."],
        ["Utilities", "server/utils/*.js", "Shared helpers for token signing, DB connection, async error wrapping, mail transport, PDF HTML, and Chrome path."],
    ], widths=[1.45, 2.0, 3.05], header_fill=LIGHT_BLUE)
    add_para(doc, "The backend follows a practical MVC-service style. Routes stay thin, controllers coordinate, models define persistence, and services perform background or external operations. This separation helps explain the code during a meeting: if someone asks where meeting creation happens, the URL is in routes/meetingRoutes.js, the business flow is in controllers/meetingController.js, and supporting integrations are in services.")

    doc.add_heading("8. Authentication and Workspace Security", level=1)
    add_image(doc, diagrams["auth"], "Figure 2: Login, token storage, user lookup, and workspace membership validation.")
    add_para(doc, "The authentication flow starts in the frontend. AuthProvider reads mom.token from localStorage, sets it on the Axios client, and calls /auth/me when a token exists. Login and register requests return a token plus sanitized user information. If any protected API later returns 401, the response interceptor clears user and token state, forcing the app back into an unauthenticated condition.")
    add_para(doc, "On the backend, requireAuth supports two authentication modes. For normal web users, it reads the Authorization header, verifies the JWT with JWT_SECRET, loads the user from MongoDB, removes passwordHash from the selected fields, and attaches req.user. For third-party or programmatic access, it can also accept x-api-key, hash it with SHA-256, look up an active ApiKey record, attach its user, and optionally inject the API key's restricted workspace.")
    add_para(doc, "Workspace security is handled separately by requireWorkspace. This middleware reads x-workspace-id from the request, or uses req.apiKeyWorkspace when an API key is workspace-restricted. It checks the Membership collection for the authenticated user and selected workspace. If the workspace id is missing, it falls back to the user's first membership. If no membership exists, the request is rejected with workspace access failure.")
    add_table(doc, ["Security Item", "Implementation Detail", "Impact"], [
        ["JWT token", "Stored as mom.token in browser and sent as Bearer token", "Keeps protected routes authenticated across refreshes."],
        ["Workspace id", "Stored as mom.workspaceId and sent as x-workspace-id", "Prevents accidental cross-workspace data access."],
        ["Membership role", "Membership role can be owner/editor/viewer", "Provides role context for future or current route restrictions."],
        ["API key", "Stored hashed; raw key only returned once", "Allows safer integration-style access without storing raw keys."],
        ["Password hash", "User model stores passwordHash, auth responses sanitize user", "Avoids exposing credential hashes to client responses."],
    ], widths=[1.45, 2.7, 2.35])

    doc.add_heading("9. Meeting Scheduling and Lifecycle", level=1)
    add_image(doc, diagrams["meeting"], "Figure 3: Meeting creation, participant invites, start/end lifecycle, MOM/PDF generation.")
    add_para(doc, "Meeting creation begins in CreateMeeting.jsx. The page renders MeetingForm and submits the payload to POST /create-meeting. The backend validates the title, duration, online/offline requirements, and start time. It calculates endTime from startTime plus duration, normalizes participants, ensures the organizer is included, and then attempts platform-specific online meeting creation when the meeting is online.")
    add_para(doc, "Participant normalization is important because the application supports both registered users and external guests. The controller checks participant emails, matches existing User records, assigns kind=user or kind=external, creates invite tokens, sets organizer status to joined, and removes duplicates. Invite tokens are used to generate public join URLs without exposing raw internal permissions.")
    add_para(doc, "If the platform is Google Meet, Teams, or Zoom, the controller calls the corresponding service. Google and Microsoft integrations can create calendar events and meeting links. Zoom service can create a Zoom meeting where configured. If an integration fails, meeting creation can still continue; the application records the meeting and leaves link behavior dependent on the available result.")
    add_para(doc, "Meeting lifecycle status moves through scheduled, ongoing, completed, or cancelled. The Start Meeting button calls POST /meeting/:id/start, which marks the meeting ongoing and records actualStartTime. The End Meeting button calls POST /meeting/:id/end, which triggers internalEndMeeting. That routine marks participants inactive, generates attendance, creates or updates MOM summary data, generates a PDF, stores pdfUrl, and saves the meeting.")
    add_table(doc, ["Lifecycle Step", "Frontend Trigger", "Backend Behavior"], [
        ["Create", "CreateMeeting -> MeetingForm submit", "Validates payload, normalizes participants, creates optional online link, saves Meeting, sends invites."],
        ["Invite", "MeetingDetails invite box", "Adds new participant, creates invite token, sends email with /join/:id?invite=token link."],
        ["Join", "JoinMeeting public page or protected join", "Validates token/user, marks participant joined, updates RSVP, joinedAt, and active status."],
        ["Start", "MeetingDetails Start Meeting button", "Changes status to ongoing and records actualStartTime."],
        ["End", "MeetingDetails End Meeting button or cron", "Generates attendance, MOM output, PDF, and marks status completed."],
        ["Cancel/Delete", "Meeting list/detail actions", "Cancel changes status; delete removes meeting plus related MOM/action items/notifications."],
    ], widths=[1.3, 2.0, 3.2], header_fill=LIGHT_BLUE)

    doc.add_heading("10. MOM Editor and Document Workflow", level=1)
    add_image(doc, diagrams["mom"], "Figure 4: MOM editor load, template application, autosave, action extraction, publish, and PDF trigger.")
    add_para(doc, "The MOM editor is centered around MinutesEditor.jsx and WordLikeEditor.jsx. MinutesEditor loads meeting and existing MOM data from GET /meeting/:id/minutes. If no existing content is available, it applies either the custom editor template from /editor-template or a default MOM template from meeting utilities. Placeholders such as [DATE], [TIME], [CREATOR], [PARTICIPANTS], [MEETING_TITLE], and [AGENDA] are replaced using the loaded meeting and user data.")
    add_para(doc, "WordLikeEditor is a TipTap-based editor that behaves like a lightweight Word document inside the browser. It supports headings, paragraph text, bold, italic, underline, color, alignment, bullet lists, ordered lists, links, images, and tables. The editor maintains dirty/saving/saved/error states and automatically saves after a delay when content changes. Manual save is also available.")
    add_para(doc, "Every save sends contentHtml to PUT /meeting/:id/minutes. The backend verifies that the user is either the meeting creator or a participant, normalizes attachment ids, chooses draft or published status, extracts generated action items from the HTML, merges them with manual action items, upserts the MOM document, increments version, saves meeting.momContent, and synchronizes ActionItem records.")
    add_para(doc, "Publishing a MOM sets docStatus to published and can trigger the end-of-meeting lifecycle when the save is manual. This means a user can author the MOM during or after the meeting and then publish it to generate final artifacts. The Download PDF button first saves the latest editor state, then opens /meeting/:id/download-pdf with regen=1 so the backend regenerates the PDF from the latest content.")
    add_table(doc, ["MOM Data Area", "Where It Appears", "Stored In"], [
        ["Rich text body", "WordLikeEditor document canvas", "Mom.contentHtml and Meeting.momContent"],
        ["Document status", "Right side panel", "Mom.docStatus: draft or published"],
        ["Summary and decisions", "Extra Details panel", "Mom.summary and Mom.decisions"],
        ["Action items", "Generated Tasks panel and ActionItems page", "Mom.actionItems and ActionItem collection"],
        ["Attachments/images", "Editor image upload and attachments list", "Attachment model and /uploads files"],
        ["Version", "Backend update counter", "Mom.version increments on each upsert"],
    ], widths=[1.6, 2.35, 2.55])

    doc.add_heading("11. Real-Time Collaboration Design", level=1)
    add_para(doc, "The project includes a dedicated collaboration service using Hocuspocus. server/services/collabService.js creates a Hocuspocus Server named mom-collaboration and attaches a Database extension. The extension reads and writes binary Yjs document state into a MongoDB collection named yjs_documents. This design separates collaboration state from the normal MOM HTML save path.")
    add_para(doc, "The collaboration server has authentication through JWT. When a collaborative client connects, onAuthenticate checks for a token and verifies it with JWT_SECRET. If valid, the collaboration context receives a user id and display name. If invalid or missing, the collaboration connection is rejected. The server also logs document load events, which helps during debugging.")
    add_para(doc, "In the current frontend, the WordLikeEditor imports TipTap and editor extensions and the package list includes Yjs, y-prosemirror, Hocuspocus provider, and collaboration extensions. The backend foundation for real-time collaborative persistence is present, and the editor architecture is compatible with extending the TipTap setup to connect each meeting document to the Hocuspocus server by documentName.")
    add_callout(doc, "Collaboration explanation", "For presentation, describe this as a two-layer save model: normal MOM saves write HTML and structured fields through REST APIs, while Hocuspocus/Yjs is designed to sync live editor changes between users and persist collaborative state in MongoDB.")

    doc.add_heading("12. Action Items and Task Tracking", level=1)
    add_para(doc, "Action item tracking connects the meeting document to post-meeting execution. The editor detects sentences like a person will do something by a date and shows generated tasks in real time. The backend also extracts action items from contentHtml using momActionItemService, merges generated tasks with manual items, and then syncs them into the ActionItem model through actionItemService.")
    add_para(doc, "The ActionItem collection is separate from Mom.actionItems so that the application can query tasks independently from full MOM documents. ActionItems.jsx fetches /action-items with optional scope, status, and search filters. Updating a task uses PATCH /action-items/:id/status and can switch a task between pending and completed. This gives employees a simple personal task tracker while still keeping the MOM as the source document.")
    add_table(doc, ["Action Item Field", "Purpose"], [
        ["workspaceId", "Keeps task inside tenant/workspace boundary."],
        ["meetingId", "Links task back to the meeting where it was created."],
        ["momId", "Links task to the specific MOM document."],
        ["sourceItemId", "Can track item identity from mom.actionItems."],
        ["title/task", "Readable task description shown in tracker."],
        ["assignedTo", "Person/email/name responsible for completing the task."],
        ["deadline", "Due date for reminder and sorting logic."],
        ["status", "pending or completed."],
    ], widths=[1.7, 4.8], header_fill=LIGHT_BLUE)

    doc.add_heading("13. Reminders, Notifications, and Cron Jobs", level=1)
    add_para(doc, "The reminder loop starts after MongoDB connects and the Express server begins listening. startReminderLoop in reminderService.js immediately runs a tick and then schedules future ticks using node-cron. The cron expression comes from REMINDER_CRON or defaults to every minute, and timezone comes from REMINDER_TIMEZONE or Asia/Kolkata.")
    add_para(doc, "The loop performs four major tasks. First, it processes meeting reminders for scheduled meetings with reminderMinutes configured and no reminderSentAt. It creates Notification records for registered users, sends email reminders, and marks reminderSentAt. Second, it checks action items due within the next 24 hours and creates due-soon notifications. Third, it auto-completes scheduled meetings whose date is older than today. Fourth, it auto-completes ongoing meetings whose endTime has passed by invoking internalEndMeeting.")
    add_table(doc, ["Cron Function", "What It Does", "Database Effect"], [
        ["processMeetingReminders", "Finds scheduled meetings whose reminder time has arrived.", "Creates Notification documents, sends emails, sets reminderSentAt."],
        ["processActionItemReminders", "Finds pending MOM action items due within 24 hours.", "Creates due-soon notifications and marks reminder sent on the item."],
        ["autoCompletePastScheduledMeetings", "Finds old scheduled meetings.", "Marks status completed."],
        ["autoCompleteOngoingMeetings", "Finds ongoing meetings where endTime passed.", "Runs internalEndMeeting to finalize attendance, MOM, and PDF."],
    ], widths=[2.1, 2.3, 2.1])
    add_para(doc, "The frontend also has a user-facing upcoming meeting check in App.jsx. When a user is authenticated and a workspace is active, the app polls GET /upcoming every minute, examines today's meetings, calculates the next future start time, and displays a toast if the meeting starts within ten minutes. This gives users immediate visual reminders even before backend notification screens are opened.")

    doc.add_heading("14. PDF Export and Sharing", level=1)
    add_para(doc, "PDF generation appears in two backend paths. The meeting lifecycle uses generatePDF in meetingController.js, which builds HTML using buildMomHtml, opens it in Puppeteer, renders an A4 PDF, saves it to server/uploads/{meetingId}.pdf, and stores meeting.pdfUrl. The download endpoint can regenerate the PDF when regen=1 is provided and then sends the file to the browser.")
    add_para(doc, "The separate pdfService and pdfController also provide a generate-pdf route. Together, these pieces show that the project treats PDF as a backend-generated artifact, not a browser print shortcut. This is important because it allows consistent formatting, server-side regeneration, and permanent file URLs under /uploads.")
    add_para(doc, "Sharing is token-based. POST /share creates a Share document containing workspaceId, meetingId, unique token, accessType, expiry, and createdBy. GET /share/:token opens the shared MOM. PATCH /share/:token/minutes can update shared minutes when accessType permits editing. Share tokens allow read-only or controlled edit access without requiring a full authenticated workspace session.")
    add_table(doc, ["Feature", "Route/Model", "Presentation Explanation"], [
        ["Download PDF", "GET /meeting/:id/download-pdf", "Saves latest content first, regenerates with Puppeteer, and opens PDF in a new browser tab."],
        ["Lifecycle PDF", "internalEndMeeting -> generatePDF", "Meeting completion automatically creates a PDF artifact."],
        ["Share link", "Share model and /share routes", "Creates a secure token with expiry and access type."],
        ["Static files", "app.use('/uploads', express.static(...))", "Generated PDFs and uploaded files become accessible through server URLs."],
    ], widths=[1.5, 2.2, 2.8], header_fill=LIGHT_BLUE)

    doc.add_heading("15. Integrations: Google, Microsoft Teams, Zoom, and AI", level=1)
    add_para(doc, "The Settings page exposes integration status, connect, disconnect, sync, preferences, and API key management. Google integration routes handle status, connect URL generation, OAuth callback, disconnect, sync, and preference update. Microsoft integration routes mirror that structure for Microsoft Graph. Zoom service is called during meeting creation when platform is zoom.")
    add_para(doc, "Google Calendar service creates Google Meet events and can store event ids and Meet links in the Meeting.graph object. Microsoft Graph service creates Outlook/Teams events, stores event ids, and supports token handling through MSAL. Zoom service creates Zoom meetings where credentials are configured. Each integration is written as a service so controller logic can remain readable.")
    add_para(doc, "AI routes include summarize, extract-tasks, and polish. aiService uses the Google Generative AI dependency when configured. In presentation terms, this gives the project an extensible layer for improving MOM content: summaries can be generated, action items can be extracted, and text can be polished without hard-coding those capabilities into the editor component.")
    add_table(doc, ["Integration", "Main Files", "Role in Workflow"], [
        ["Google", "googleController.js, googleCalendarService.js", "OAuth connect, status, sync, Google Meet event creation, preferences."],
        ["Microsoft/Teams", "integrationController.js, microsoftGraphService.js", "OAuth connect/callback, Graph event creation, Teams meeting links, sync preferences."],
        ["Zoom", "zoomService.js", "Creates Zoom meetings and optional host link during meeting creation."],
        ["AI", "aiRoutes.js, aiController.js, aiService.js", "Summarize, extract tasks, and polish MOM text."],
        ["Email", "emailService.js, mailer.js", "Sends meeting invitations, reminders, and post-meeting style communications."],
    ], widths=[1.4, 2.1, 3.0])

    doc.add_heading("16. Database Schema", level=1)
    add_image(doc, diagrams["data"], "Figure 5: Main MongoDB collections and relationships.")
    add_para(doc, "The database is MongoDB with Mongoose models. Workspace and Membership create the tenant boundary. User stores authentication identity. Meeting stores scheduling and participant lifecycle data. Mom stores the meeting minutes document and rich content. ActionItem stores follow-up tasks separately for fast tracking. Share stores public access tokens. Notification stores reminder and due-soon records. Attachment stores files referenced by MOM or meetings. IntegrationToken and ApiKey support external service access.")
    add_table(doc, ["Model", "Key Fields", "Explanation"], [
        ["User", "email, name, passwordHash", "Authenticated application user. Password hash is never returned in normal auth responses."],
        ["Workspace", "name, createdBy", "Tenant/project container for data separation."],
        ["Membership", "workspaceId, userId, role", "Maps users to workspaces with owner/editor/viewer roles."],
        ["Meeting", "workspaceId, title, agenda, date, startTime, endTime, participants, status, graph, pdfUrl", "Core meeting record and lifecycle state."],
        ["Mom", "meetingId, contentHtml, docStatus, version, actionItems, attendees, discussion, approvals", "Complete minutes document linked to one meeting."],
        ["ActionItem", "workspaceId, meetingId, momId, task, assignedTo, deadline, status", "Standalone task tracker derived from MOM content."],
        ["Share", "meetingId, token, accessType, expiry", "Public or limited MOM sharing link."],
        ["Notification", "type, title, message, entityType, dueAt", "In-app reminder and task due notifications."],
        ["Attachment", "file metadata, entityType/entityId", "Uploaded files and images referenced by MOM or meetings."],
    ], widths=[1.25, 2.5, 2.75], header_fill=LIGHT_BLUE)

    doc.add_heading("17. API Endpoint Reference", level=1)
    add_para(doc, "All main APIs are mounted under /api/v1. The frontend API client already includes this prefix, so components call shorter paths such as /meetings or /meeting/:id/minutes. Most project APIs require requireAuth and requireWorkspace; public join and share routes are exceptions where token-based access is used.")
    add_table(doc, ["Area", "Endpoints", "Purpose"], [
        ["Auth", "POST /auth/register, POST /auth/login, GET /auth/me, POST /auth/verify-documents", "Register/login user, validate session, and document verification placeholder flow."],
        ["Workspace", "GET /workspaces, POST /workspaces", "List user's workspace memberships and create workspace."],
        ["Meetings", "POST /create-meeting, GET /meetings, GET/PATCH/DELETE /meeting/:id", "Create, list, view, update, and delete meetings."],
        ["Lifecycle", "POST /meeting/:id/start, POST /meeting/:id/end, POST /meeting/:id/join, POST /meeting/:id/leave", "Start/end meetings and track participant attendance."],
        ["Invites", "POST /meeting/:id/invite, GET /join/:id, POST /join/:id/accept", "Invite additional participants and allow token-based public acceptance."],
        ["MOM", "POST /create-mom, GET /mom/:meetingId, GET/PUT /meeting/:meetingId/minutes", "Create/read/update MOM document and rich editor content."],
        ["Action Items", "GET /action-items, PATCH /action-items/:id/status", "List and update task status."],
        ["Attachments", "POST /attachments, GET /attachments", "Upload and list files/images."],
        ["Notifications", "GET /notifications, POST /notifications/:id/read", "Read and mark notifications."],
        ["Share", "POST /share, GET /share/:token, PATCH /share/:token/minutes", "Create and open secure share links."],
        ["PDF/Reports", "GET /meeting/:id/download-pdf, GET /generate-pdf/:meetingId, GET /api/reports/summary", "Generate/download PDFs and fetch report summary."],
        ["Integrations", "Google/Microsoft status/connect/callback/disconnect/sync/preferences", "Manage calendar and online meeting integrations."],
        ["AI/API keys", "POST /ai/*, POST/GET/DELETE /apikeys", "Summarize/extract/polish content and manage API keys."],
    ], widths=[1.25, 2.75, 2.5])

    doc.add_heading("18. Error Handling, Validation, and Security", level=1)
    add_para(doc, "The backend uses asyncHandler to wrap async controller functions and forward errors to Express error middleware. notFound handles unmatched routes, and errorHandler returns structured errors. Many controllers set HTTP status before throwing an Error, which gives the client a meaningful response code and message. The frontend displays errors through toast notifications or page-level fallback states.")
    add_para(doc, "Validation is a mix of explicit controller checks and schema-level restrictions. For example, createMeeting checks title, duration, start time, offline location, and online platform. Mongoose enums restrict meeting status, participant RSVP, MOM doc status, action item status, and other fields. Workspace middleware guards tenant access before protected controllers perform queries.")
    add_para(doc, "Security concerns are addressed through hashed passwords, JWT verification, API key hashing, workspace membership checks, invite tokens, share token expiry, and sanitized participant responses. The backend deletes inviteToken before returning participants through normal meeting APIs, reducing accidental exposure of public join credentials.")
    add_callout(doc, "Important observation", "Some visitor-related frontend/backend files exist, but key visitor route mounting is commented out in server.js and several visitor routes in App.jsx are inside comments. For the main MOM documentation, visitor management should be presented as an adjacent or partially integrated module, not as the primary active workflow.")

    doc.add_heading("19. Deployment and Environment Configuration", level=1)
    add_para(doc, "The project has separate client and server package.json files. The client runs with Vite using npm run dev, and the server runs with node server.js or nodemon through npm run dev. Environment variables are stored separately in client/.env and server/.env with examples available. The frontend can use VITE_API_BASE_URL; otherwise, it detects the browser hostname and points to port 5000 for the backend.")
    add_table(doc, ["Environment Area", "Variables / Config", "Purpose"], [
        ["Client API base", "VITE_API_BASE_URL", "Overrides automatic backend URL detection."],
        ["Server port", "PORT / loadConfig", "Controls Express server port, typically 5000."],
        ["MongoDB", "MONGO_URI", "Database connection string used by connectDB."],
        ["Auth", "JWT_SECRET", "Signs and verifies JWT access tokens."],
        ["Public URLs", "PUBLIC_CLIENT_BASE_URL, PUBLIC_API_BASE_URL", "Used for invite URLs and PDF asset base paths."],
        ["Reminders", "REMINDER_CRON, REMINDER_TIMEZONE", "Controls cron frequency and timezone."],
        ["Integrations", "Google, Microsoft, Zoom credentials", "Allow calendar/video meeting creation and sync."],
        ["Email", "SMTP-related variables", "Used by nodemailer/email service for invites and reminders."],
        ["Collaboration", "COLLAB_PORT or config equivalent", "Starts Hocuspocus collaboration server separately from REST API."],
    ], widths=[1.6, 2.15, 2.75], header_fill=LIGHT_BLUE)
    add_para(doc, "For deployment, the backend should run as a long-lived Node process because cron reminders and collaboration server startup depend on the server process. The uploads directory must be writable and persistent if generated PDFs and attachments should survive restarts. MongoDB must be reachable from the backend. The frontend build can be hosted separately as static files as long as its API base URL points to the deployed backend.")

    doc.add_heading("20. End-to-End Workflow Summary", level=1)
    add_numbered(doc, [
        "User opens the app and either registers or logs in. AuthProvider stores the token and verifies the user through /auth/me.",
        "WorkspaceProvider loads workspace memberships and selects an active workspace. Axios adds x-workspace-id to protected API requests.",
        "User creates a meeting from the CreateMeeting page. The backend validates the schedule, normalizes participants, attempts online platform link creation, saves the Meeting, and sends invitations.",
        "Participants receive tokenized join links. Internal users can access meeting screens; external users can accept the invite through the public join route.",
        "Before the meeting, reminder logic can send notifications/emails, and the frontend can show a toast when a meeting starts within ten minutes.",
        "The organizer starts the meeting, which changes status to ongoing and records actualStartTime.",
        "During or after the meeting, the user opens MinutesEditor. It loads meeting context and existing MOM content, applies a template if needed, and lets the user write rich formatted minutes.",
        "Autosave or manual save sends contentHtml and action items to the backend. The MOM is upserted, version is incremented, action items are extracted/merged, and the ActionItem collection is synchronized.",
        "When the MOM is published or the meeting is ended, backend lifecycle logic generates attendance, MOM summary data, and a PDF artifact.",
        "The user can download the PDF, share a secure token link, view reports, track action items, and mark assigned tasks completed.",
    ])
    add_callout(doc, "Presentation close", "The strongest way to present the project is as one continuous loop: Plan the meeting -> Run the meeting -> Record decisions -> Generate tasks -> Share final MOM -> Track completion.")

    doc.add_heading("21. Future Improvements", level=1)
    add_table(doc, ["Improvement", "Reason", "Suggested Approach"], [
        ["Complete live editor collaboration UI", "Backend Hocuspocus service exists, but frontend can be expanded to connect editor sessions by meeting id.", "Add HocuspocusProvider and TipTap Collaboration extension in WordLikeEditor with documentName based on meeting id."],
        ["Stronger role permissions", "Membership roles exist but not every controller enforces role-specific actions.", "Use requireWorkspaceRole on create/update/delete/publish endpoints as needed."],
        ["Better audit trail", "MOM version increments but detailed change history is not exposed.", "Create a MomVersion or ActivityLog model for saves, publishes, downloads, and share events."],
        ["Improved validation schema", "Some validation is manual in controllers.", "Use zod schemas with validateMiddleware for consistent payload validation."],
        ["Visitor module cleanup", "Visitor files exist but route mounting is partly disabled.", "Either fully integrate visitor routes or move them into a clearly separate module."],
        ["Production file storage", "Uploads are local to server filesystem.", "Use cloud object storage for PDFs and attachments in deployed environments."],
        ["Testing", "Current repo has limited formal test coverage.", "Add API tests for auth/workspace, meeting lifecycle, MOM save, action-item sync, and share/PDF routes."],
    ], widths=[1.7, 2.45, 2.35], header_fill=LIGHT_BLUE)
    add_para(doc, "These improvements do not reduce the current project's value. They are practical next steps that make the system more production-ready, easier to maintain, and easier to demonstrate as a professional project in future reviews.")

    add_detailed_appendices(doc)


def add_detailed_appendices(doc):
    doc.add_page_break()
    doc.add_heading("Appendix A. Frontend Page-by-Page Walkthrough", level=1)
    add_para(doc, "This appendix explains how each major frontend screen contributes to the user journey. It is useful for a meeting presentation because it lets the presenter move screen by screen and connect the UI to the backend APIs behind it. The frontend is not just a collection of React files; it is organized around the daily workflow of a meeting owner, participant, employee, and admin-style user.")
    add_table(doc, ["Screen / Component", "User Role in Flow", "Implementation Detail"], [
        ["Login.jsx", "Entry point for existing users.", "Collects email/password and calls AuthProvider.login, which posts to /auth/login and stores the returned JWT."],
        ["Register.jsx", "Entry point for new users.", "Calls AuthProvider.register, which posts name, email, and password to /auth/register and initializes the session."],
        ["ProtectedRoute.jsx", "Security boundary for authenticated pages.", "Prevents unauthenticated users from entering workspace screens before AuthProvider has a valid user."],
        ["AppLayout.jsx", "Main shell for protected screens.", "Provides the consistent application navigation and layout around dashboard, meetings, calendar, MOM, settings, reports, and action items."],
        ["Dashboard.jsx", "Quick operational overview.", "Fetches meetings and MOM snippets to show recent/upcoming activity and provide fast navigation into details."],
        ["Meetings.jsx", "Meeting management list.", "Fetches /meetings, displays meeting cards, supports update/delete mutations, and routes users to meeting details."],
        ["MeetingCard.jsx", "Reusable meeting preview.", "Shows title, schedule, status, participants, and action shortcuts in a compact card format."],
        ["CreateMeeting.jsx", "Meeting creation workflow.", "Wraps MeetingForm and posts the final payload to /create-meeting through React Query mutation."],
        ["MeetingForm.jsx", "Meeting input engine.", "Collects title, agenda, date/time, duration, participants, priority, type, platform, location, reminders, and recurring-style flags."],
        ["ParticipantPicker.jsx", "Participant discovery.", "Calls /users and allows internal users or email-style participants to be selected for the meeting."],
        ["MeetingDetails.jsx", "Main meeting control room.", "Loads /meeting/:id and /mom/:id, starts/ends meetings, invites participants, previews MOM content, and links to PDF and editor."],
        ["JoinMeeting.jsx", "External invite acceptance.", "Reads invite token from URL, calls /join/:id and /join/:id/accept, then shows safe meeting link data."],
        ["Calendar.jsx + CalendarView.jsx", "Date-based meeting view.", "Fetches meetings and renders them through FullCalendar components for schedule awareness."],
        ["MinutesEditor.jsx", "MOM authoring workflow.", "Loads meeting/MOM, applies templates, manages save/publish/export, and coordinates editor content with action items."],
        ["WordLikeEditor.jsx", "Rich-text document editor.", "TipTap editor with toolbar controls, autosave, image upload, tables, formatting, task detection, and document-style canvas."],
        ["ActionItems.jsx", "Post-meeting task tracker.", "Calls /action-items with filters and updates task status through PATCH endpoint."],
        ["Documents.jsx", "Document-oriented meeting output.", "Uses meeting data to organize generated or attached documents for retrieval."],
        ["Reports.jsx", "Management summary.", "Calls /api/reports/summary for higher-level project/meeting statistics."],
        ["Settings.jsx", "Integration and key management.", "Controls Google/Microsoft integration flows, sync preferences, connect/disconnect, and API key creation/revocation."],
        ["TemplateBuilder.jsx", "MOM template customization.", "Loads and saves editor template HTML through /editor-template so new MOM documents can start with a chosen format."],
        ["ShareView.jsx", "Token-based shared MOM view.", "Loads /share/:token and can PATCH shared minutes when the token allows edit access."],
        ["Notifications.jsx", "Reminder inbox.", "Fetches /notifications and lets the user mark notifications as read."],
        ["Profile.jsx", "User activity/profile area.", "Fetches meetings and action items for profile-style context."],
        ["AdminDashboard.jsx / EmployeeDashboard.jsx", "Role-flavored dashboards.", "Provide alternate dashboard views for different responsibilities using the same backend data foundation."],
    ], widths=[1.55, 1.75, 3.2], header_fill=LIGHT_BLUE)
    add_para(doc, "A presenter can describe the frontend as four connected zones. The first zone is access management: login, register, protected route, and workspace provider. The second zone is meeting management: dashboard, meeting list, calendar, create form, details, join page, and participant picker. The third zone is documentation: MOM editor, template builder, documents, share view, attachments, and PDF export. The fourth zone is follow-up and administration: action items, notifications, reports, settings, integrations, and profile.")
    add_para(doc, "The frontend also uses a consistent mutation pattern. Whenever a user creates a meeting, updates a meeting, saves MOM content, marks an action item, or changes integration settings, the relevant React Query cache is invalidated. This makes the interface feel connected: a MOM save refreshes the action item tracker, a meeting update refreshes meeting cards, and publishing minutes refreshes meeting detail data.")
    add_para(doc, "State management is intentionally lightweight. There is no large external global store. Authentication and workspace are handled by React context providers. Server data is handled by React Query. Local screen state is handled with useState and useMemo. This makes the project easier to explain because each state type has a clear owner: session state in providers, server state in queries, and form/editor state inside components.")

    doc.add_heading("Appendix B. Backend Controller and Service Deep Dive", level=1)
    add_para(doc, "This appendix connects each backend module to its responsibility. The backend is where the project becomes more than a UI: it protects data, normalizes participants, talks to external services, stores records, creates PDFs, schedules reminders, and keeps tasks synchronized.")
    add_table(doc, ["Controller / Service", "Primary Responsibilities", "Important Implementation Notes"], [
        ["authController + authService", "Register, login, and current-user loading.", "authService hashes passwords with bcryptjs, compares login credentials, signs JWTs, and returns sanitized user objects."],
        ["workspaceController", "List and create workspaces.", "When a workspace is created, membership can associate the user with a role, allowing workspace-aware data separation."],
        ["meetingController", "Largest business controller for meetings.", "Handles create/list/detail/update/delete, invite, join/leave, start/end, agenda update, public invite acceptance, PDF download, and lifecycle generation."],
        ["momController", "Create, load, and upsert MOM documents.", "Verifies meeting access, saves rich HTML, handles publish status, merges action items, syncs ActionItem records, and updates meeting.momContent."],
        ["actionItemController", "Task tracker API.", "Lists action items using workspace context and changes pending/completed status."],
        ["attachmentController", "File upload/list flow.", "Uses multer-style upload handling, saves files under uploads, and records metadata for MOM or meeting entities."],
        ["shareController", "Secure share links.", "Creates share tokens, opens shared MOM documents, and supports token-based update depending on access type."],
        ["notificationController", "Notification list and read state.", "Allows frontend notification inbox to show reminders and mark them read."],
        ["reportController", "Summary metrics.", "Provides dashboard/reporting style aggregate data for management presentation."],
        ["editorTemplateController", "MOM template persistence.", "Gets or upserts the workspace's editor template used by MinutesEditor defaults."],
        ["integrationController", "Microsoft integration endpoints.", "Wraps Microsoft Graph connect/callback/status/disconnect/sync/preference behavior."],
        ["googleController", "Google integration endpoints.", "Wraps Google OAuth and calendar/Meet behavior."],
        ["apiKeyController", "API key lifecycle.", "Creates hashed keys, lists metadata, and revokes active keys."],
        ["aiController + aiService", "AI helper endpoints.", "Summarizes content, extracts tasks, and polishes MOM text when provider configuration is present."],
        ["reminderService", "Scheduled background operations.", "Runs meeting reminders, action item due reminders, old meeting completion, and ongoing meeting auto-completion."],
        ["pdfService + pdfTemplate", "MOM PDF rendering.", "Builds HTML and renders PDF through Puppeteer/Chrome path resolution."],
        ["emailService + mailer", "Email delivery.", "Sends invitations, reminders, and related communication using configured mail transport."],
        ["collabService", "Real-time document sync foundation.", "Starts Hocuspocus server and persists Yjs binary state in MongoDB yjs_documents collection."],
        ["googleCalendarService", "Google Calendar/Meet events.", "Creates/updates/deletes Google event records and manages OAuth client behavior."],
        ["microsoftGraphService", "Microsoft Teams/Outlook events.", "Uses MSAL/Graph to create online meeting events and maintain token cache state."],
        ["zoomService", "Zoom meeting creation.", "Requests Zoom access token and creates meeting links where credentials are configured."],
        ["momActionItemService", "Task extraction logic.", "Parses content text for action-style sentences, normalizes dates, title-cases assignees, and merges generated/manual tasks."],
        ["actionItemService", "Task synchronization.", "Copies MOM action items into the ActionItem collection so they can be tracked independently."],
    ], widths=[1.65, 2.1, 2.75])
    add_para(doc, "meetingController deserves special attention because it contains the complete meeting lifecycle. It converts raw participants into normalized participant subdocuments, protects invite tokens, creates online meeting links through services, sends invitations in the background, supports public acceptance, and finalizes meetings through internalEndMeeting. This makes it the best backend file to explain when presenting the project's core logic.")
    add_para(doc, "momController is the second most important backend file for the project's identity. It connects the document editor to persistence. It ensures the meeting exists inside the active workspace, checks whether the editor is creator or participant, saves contentHtml, increments version, manages draft/published status, extracts action items, populates attachments, updates the meeting preview content, and can trigger report generation after manual publish.")
    add_para(doc, "Services reduce controller complexity. Instead of putting email formatting, PDF rendering, calendar OAuth, action item parsing, and cron loops directly into controllers, those responsibilities live in service modules. This is good architecture for a student or professional project because it shows separation of concerns and makes future testing easier.")

    doc.add_heading("Appendix C. Detailed Workflow Narratives", level=1)
    add_para(doc, "The following narratives explain the most important flows in complete sentences. These can be used directly while speaking in a presentation because they translate code paths into human-readable steps.")
    doc.add_heading("C1. User Login to Dashboard", level=2)
    add_numbered(doc, [
        "The user enters email and password on the Login page.",
        "AuthProvider.login sends POST /auth/login with credentials.",
        "authController passes the request to authService.login.",
        "authService finds the user by email, compares password with bcryptjs, signs a JWT, and returns token plus sanitized user.",
        "AuthProvider stores the token in mom.token and sets the Authorization header on the Axios instance.",
        "WorkspaceProvider detects that the user is authenticated and calls GET /workspaces.",
        "The backend verifies JWT, loads memberships, and returns workspace objects with roles.",
        "The frontend selects an active workspace and stores it as mom.workspaceId.",
        "ProtectedRoute allows access to AppLayout, and the user lands on Dashboard or the requested protected route.",
    ])
    doc.add_heading("C2. Meeting Creation with Participants", level=2)
    add_numbered(doc, [
        "The user opens Create Meeting and fills title, agenda, schedule, duration, meeting type, platform/location, participants, and reminder settings.",
        "MeetingForm returns a payload to CreateMeeting.jsx.",
        "CreateMeeting posts the payload to /create-meeting.",
        "requireAuth verifies the JWT and requireWorkspace resolves workspace membership.",
        "createMeeting validates title, duration, start time, online platform, or offline location.",
        "Participants are normalized: duplicate emails are removed, registered users are linked to User ids, external participants are marked external, and invite tokens are created.",
        "If the meeting is online, the controller calls Google, Microsoft, or Zoom service depending on the selected platform.",
        "The Meeting document is saved with workspaceId, schedule, status scheduled, participant data, reminder settings, platform graph data, and createdBy.",
        "Background invitation email sending begins for invited participants, and the frontend navigates back to the meeting list after success.",
    ])
    doc.add_heading("C3. MOM Drafting and Publishing", level=2)
    add_numbered(doc, [
        "The user opens Meeting Details and clicks Open MOM editor.",
        "MinutesEditor loads meeting and MOM data through GET /meeting/:id/minutes.",
        "If content is blank, the editor uses the workspace template or default MOM template and replaces placeholders with meeting data.",
        "WordLikeEditor lets the user write rich content, insert images, create tables, format text, and use headings/lists.",
        "On content change, the editor updates local HTML state, detects action-item style sentences, and starts autosave timing.",
        "Autosave or manual save calls PUT /meeting/:id/minutes with contentHtml, docStatus, actionItems, summary, and decisions.",
        "momController checks permission, extracts generated action items from HTML, merges them with manual items, upserts the MOM, increments version, and syncs ActionItem records.",
        "When the user publishes, docStatus becomes published. If manual publish is used, internalEndMeeting can run in the background to generate final report artifacts.",
        "The Download PDF action saves latest content and opens the PDF download endpoint with regeneration enabled.",
    ])
    doc.add_heading("C4. Action Item Lifecycle", level=2)
    add_numbered(doc, [
        "Action items can be typed manually in the side panel or written naturally inside the MOM content.",
        "The editor detects simple will-do sentences and displays real-time generated tasks.",
        "The backend repeats extraction from contentHtml to avoid relying only on client-side detection.",
        "Generated tasks and manual tasks are merged so duplicates are minimized.",
        "syncActionItemsFromMom writes task records into the ActionItem collection with workspace, meeting, mom, assignee, deadline, and status.",
        "Employees or users open ActionItems.jsx and filter tasks by status or search.",
        "When a task is completed, the frontend patches the action item status and invalidates task queries.",
        "Reminder service can detect due soon tasks and create notifications for assigned registered users.",
    ])
    doc.add_heading("C5. End Meeting and Final Deliverables", level=2)
    add_numbered(doc, [
        "The meeting owner clicks End meeting and generate MOM, or the cron loop auto-completes an overdue ongoing meeting.",
        "internalEndMeeting loads the Meeting document and exits early if already completed and processed.",
        "The meeting status becomes completed and actualEndTime is recorded.",
        "Participants marked active are closed out with leftAt and lastActiveAt timestamps.",
        "Attendance is generated from participant joinedAt values.",
        "MOM summary/report data is generated through autoMomService where possible, with fallback values if generation fails.",
        "Puppeteer renders the MOM HTML template into a PDF and writes it under server/uploads.",
        "meeting.pdfUrl is updated, and future download requests can serve the generated PDF or regenerate it when requested.",
    ])

    doc.add_heading("Appendix D. Presentation Talking Points", level=1)
    add_para(doc, "For the final meeting presentation, the most effective sequence is to start with the problem, then show the user workflow, then explain the architecture, and finally show the technical depth. The following talking points can be used as speaker notes.")
    add_bullets(doc, [
        "This project solves the complete meeting documentation cycle, not just note writing.",
        "The frontend is a React SPA with protected routing, workspace-aware API calls, and a Word-like editor.",
        "The backend is Express with controllers and services, using MongoDB/Mongoose for persistence.",
        "Workspace membership is the core security boundary, and every protected business route depends on it.",
        "Meeting creation is advanced because it handles participants, invite tokens, online/offline mode, reminders, and optional integrations.",
        "The MOM editor stores rich HTML content, supports autosave, extracts action items, and can publish final minutes.",
        "Action items are separated into their own collection so follow-up work can be tracked independently from the document.",
        "Cron jobs make the system proactive: reminders are sent, due items are flagged, and meetings can be auto-completed.",
        "PDF generation is backend-controlled through Puppeteer, making final outputs consistent and shareable.",
        "The design is extensible: AI, Google, Microsoft, Zoom, Hocuspocus collaboration, API keys, and templates are already represented as separate modules.",
    ])
    add_callout(doc, "Final message", "The MOM project demonstrates a full-stack workflow application with authentication, authorization, scheduling, collaboration-ready editing, task tracking, automation, integrations, and document export.")


def main():
    diagrams = build_diagrams()
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_toc(doc)
    add_sections(doc, diagrams)
    doc.core_properties.title = "Minutes of Meeting (MOM) Complete Project Documentation"
    doc.core_properties.subject = "Project workflow, frontend, backend, file structure, diagrams, and technical documentation"
    doc.core_properties.author = "Codex"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
